
import json
import random
import os
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
import tiktoken
import argparse

# Carrega variáveis do .env
load_dotenv()

# Configuração do cliente OpenAI
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)
model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # Default to gpt-4o-mini if not specified

# Verificação das variáveis obrigatórias
openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    raise ValueError("A variável OPENAI_API_KEY não está definida no .env")

print(f"Usando modelo: {model_name}")

# Carrega o ficheiro JSON com a árvore de classificações
with open("Classification Request.json", "r", encoding="utf-8") as f:
    classification_tree = json.load(f)

# Extrai classes possíveis
def extract_classifiable_nodes(node, result=None):
    if result is None:
        result = []
    if node.get("CanClassifyDocuments"):
        result.append({
            "FullName": node["Key"]["FullName"],
            "Description": node.get("Description", ""),
            "Name": node.get("Name", ""),
            "Notes": node.get("Notes", ""),
            "IndexTerms": node.get("IndexTerms", "")
        })
    for child in node.get("Children", []):
        extract_classifiable_nodes(child, result)
    return result

classifiable_classes = extract_classifiable_nodes(classification_tree["Result"][0])

# Contador de tokens de input
def count_input_tokens(text: str, model: str) -> int:
    """Conta tokens de input para o modelo especificado.

    Tenta usar a codificação do modelo; se não existir, recorre a 'cl100k_base'.
    """
    try:
        enc = tiktoken.encoding_for_model(model)
    except Exception:
        enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))

# Função que invoca o Azure GPT para gerar conteúdo
def generate_content_with_gpt(class_data):
    description = class_data.get("Description", "").strip()
    notes = class_data.get("Notes", "").strip()
    index_terms = class_data.get("IndexTerms", "").replace("char(10)", "\n").strip()

    prompt = f"""
Estás a redigir um exemplo de documento administrativo, coerente e plausível, para uma classificação arquivística pública.

A classe a usar é:

**Classe:** {class_data['FullName']} – {class_data['Name']}
**Descrição:** {description}
**Notas:** {notes}
**Termos de Índice:** {index_terms}

Gera um documento fictício que se enquadre nesta classe. Deve seguir um estilo administrativo ou jurídico-formal. Usa um formato com cabeçalho típico de documentos oficiais (ex: DE, ASSUNTO, DATA, INFORMAÇÃO, DESPACHO ou outros campos comuns).

O conteúdo deve ser completo, bem estruturado e parecer realista, como se fosse mesmo usado numa entidade pública.
Nunca referencies o nome da classe no conteudo do ficheiro.
"""

    # Calcula tokens de input para monitorização
    input_tokens = count_input_tokens(prompt, model_name)
    print(f"Tokens de input estimados: {input_tokens}")

    try:
        response = client.chat.completions.create(
            model=model_name,  # Modelo OpenAI (ex: gpt-4o-mini, gpt-4, gpt-3.5-turbo)
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1000
        )
    except Exception as e:
        raise Exception(f"Erro ao chamar OpenAI: {str(e)}\nVerifique se o modelo '{model_name}' está disponível e se a chave API é válida.")

    # Preferir contagem real devolvida pela API, se disponível
    try:
        api_prompt_tokens = getattr(response, "usage", None).prompt_tokens  # type: ignore[attr-defined]
        api_completion_tokens = getattr(response, "usage", None).completion_tokens  # type: ignore[attr-defined]
        api_total_tokens = getattr(response, "usage", None).total_tokens  # type: ignore[attr-defined]
        if api_prompt_tokens is not None:
            print(f"Tokens (API): prompt={api_prompt_tokens}, completion={api_completion_tokens}, total={api_total_tokens}")
    except Exception:
        # Se não houver usage, já imprimimos estimativa acima
        pass

    return response.choices[0].message.content or ""

# Geração dos documentos
def generate_documents(n=5, output_dir="documentos_gerados"):
    os.makedirs(output_dir, exist_ok=True)

    for i in range(1, n + 1):
        cls = random.choice(classifiable_classes)
        print(f"[{i}/{n}] A gerar documento para a classe {cls['FullName']} – {cls['Name']}")

        try:
            content = generate_content_with_gpt(cls)
        except Exception as e:
            print(f"Erro ao contactar o GPT: {e}")
            continue

        # Criar .docx
        doc = Document()
        doc.add_heading(f'Documento Gerado #{i}', 0)
        doc.add_paragraph(content)

        filename = os.path.join(output_dir, f"documento_{i}_{cls['FullName'].replace('/', '-')}.docx")
        doc.save(filename)
        print(f"✔ Guardado: {filename}")

def _parse_args():
    parser = argparse.ArgumentParser(description="Gerador de documentos com OpenAI")
    parser.add_argument("--n", type=int, default=1, help="Número de documentos a gerar (default: 1)")
    parser.add_argument("--output-dir", type=str, default="documentos_gerados", help="Diretório de saída")
    return parser.parse_args()

if __name__ == "__main__":
    args = _parse_args()
    generate_documents(n=args.n, output_dir=args.output_dir)
